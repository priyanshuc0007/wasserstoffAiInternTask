# streamlit_app/app.py
from utils import data_mapping, postprocessing, preprocessing
from models.segmentation_model import load_segmentation_model, segment_image
from models.text_extraction_model import extract_text_with_otsu
from models.summarization_model import summarize_text
import streamlit as st
import os
from PIL import Image
import docx
from docx import Document
import json
import pandas as pd
import matplotlib.pyplot as plt

# Set up paths
DATA_DIR = "data"
INPUT_DIR = os.path.join(DATA_DIR, "input_images")
OUTPUT_DIR = os.path.join(DATA_DIR, "output")
SEGMENTED_DIR = os.path.join(OUTPUT_DIR, "segmented_objects")
TEXT_EXTRACTION_DIR = os.path.join(OUTPUT_DIR, "text_extraction")
SUMMARIZATION_DIR = os.path.join(OUTPUT_DIR, "summarization")
REPORT_DIR = os.path.join(OUTPUT_DIR, "report")

# Create directories if they don't exist
os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(SEGMENTED_DIR, exist_ok=True)
os.makedirs(SUMMARIZATION_DIR, exist_ok=True)
os.makedirs(TEXT_EXTRACTION_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)

st.set_page_config(page_title="AI Pipeline for Image Segmentation", layout="wide")
st.title("AI Pipeline for Image Segmentation and Object Analysis")
st.markdown("""
    <style>
    .main {
        background-color: #f0f0f5;
    }
    </style>
    """, unsafe_allow_html=True)

# Sidebar for navigation
st.sidebar.title("Navigation")
page = st.sidebar.radio("Select a page:", ["Upload Image", "Segment and Identify", "Text Extraction", "Summarization", "Final Report"])

def upload_image():
    st.header("Upload an Image")
    uploaded_file = st.file_uploader("Choose an image...", type=["png", "jpg", "jpeg"])
    if uploaded_file is not None:
        try:
            image = Image.open(uploaded_file)
            st.image(image, caption='Uploaded Image', use_column_width=True)
            image.save(os.path.join(INPUT_DIR, uploaded_file.name))
            st.success(f"Image saved as {uploaded_file.name} in {INPUT_DIR}.")
        except Exception as e:
            st.error(f"Error saving image: {e}")

def run_segmentation_and_identification():
    st.header("Segmentation and Object Identification")
    
    # Check if there's an uploaded image
    if not os.listdir(INPUT_DIR):
        st.warning("Please upload an image first!")
        return

    image_path = os.path.join(INPUT_DIR, os.listdir(INPUT_DIR)[-1])
    image = Image.open(image_path)
    st.image(image, caption="Original Image", use_column_width=True)

    if st.button("Segment and Identify"):
        with st.spinner("Loading model and processing..."):
            try:
                image_tensor = preprocessing.preprocess_image(image)
                model_seg = load_segmentation_model()
                scores, boxes, classes = segment_image(model_seg, image_tensor)
                segmented_images = postprocessing.save_segmented_objects(image, boxes, classes, SEGMENTED_DIR)

                coco_labels_file_path = r"E:\wasserstoffAiInternTask\coco_labels.txt"
                with open(coco_labels_file_path, "r") as file:
                    coco_labels = file.read().splitlines()

                for idx, (segmented_image_path, class_idx, score) in enumerate(zip(segmented_images, classes, scores)):
                    segmented_image = Image.open(segmented_image_path)
                    label_name = coco_labels[int(class_idx) - 1]
                    st.image(segmented_image, caption=f"Segment {idx + 1}: {label_name} ({score * 100:.2f}%)", use_column_width=True)

                st.success("Segmentation and object identification completed.")
            except Exception as e:
                st.error(f"Error during segmentation: {e}")

def run_text_extraction():
    st.header("Text Extraction from Segmented Images")
    segmented_images = [img for img in os.listdir(SEGMENTED_DIR) if img.endswith((".png", ".jpg", ".jpeg"))]

    text_extraction_path = os.path.join(TEXT_EXTRACTION_DIR, "text_extraction.doc")
    if not segmented_images:
        st.warning("No segmented images found. Please run segmentation first.")
        return

    # Create the document if it doesn't exist
    if not os.path.exists(text_extraction_path):
        doc = Document()
        doc.add_heading('Text Extraction', 0)
        doc.save(text_extraction_path)

    if "extracted_texts" not in st.session_state:
        st.session_state.extracted_texts = {}

    for image_name in segmented_images:
        image_path = os.path.join(SEGMENTED_DIR, image_name)
        image = Image.open(image_path)
        st.image(image, caption=image_name, use_column_width=True)

        if st.button(f"Extract Text for {image_name}"):
            with st.spinner(f"Extracting text for {image_name}..."):
                try:
                    extracted_text = extract_text_with_otsu(image_path)

                    doc = Document(text_extraction_path)
                    doc.add_picture(image_path)
                    doc.add_paragraph(f"Image Name: {image_name}")
                    doc.add_paragraph(f"Extracted Text:\n{extracted_text}")
                    doc.save(text_extraction_path)

                    st.session_state.extracted_texts[image_name] = extracted_text
                    st.success(f"Text extracted and appended to {text_extraction_path}.")
                except Exception as e:
                    st.error(f"Error during text extraction for {image_name}: {e}")

    st.write("Text extraction completed.")
    return st.session_state.extracted_texts

def run_summarization():
    st.header("Image Summarization")
    segmented_images = os.listdir(SEGMENTED_DIR)

    if not segmented_images:
        st.warning("No segmented images found! Please run segmentation first.")
        return

    doc_file_path = os.path.join(SUMMARIZATION_DIR, "image_summarization.doc")
    if not os.path.exists(doc_file_path):
        doc = Document()
        doc.add_heading('Image Summarization', 0)
        doc.save(doc_file_path)

    if "summaries" not in st.session_state:
        st.session_state.summaries = {}

    for image_file in segmented_images:
        image_path = os.path.join(SEGMENTED_DIR, image_file)
        image = Image.open(image_path)
        st.image(image, caption=f"Segmented Image: {image_file}", use_column_width=True)

        if st.button(f"Run Summarization for {image_file}"):
            with st.spinner(f"Summarizing {image_file}..."):
                try:
                    description = summarize_text(image_path)
                    doc = Document(doc_file_path)
                    doc.add_picture(image_path)
                    doc.add_paragraph(f"Image Name: {os.path.basename(image_path)}")
                    doc.add_paragraph(f"Description: {description}")
                    doc.save(doc_file_path)

                    st.session_state.summaries[image_file] = description
                    st.success(f"Summarization complete for {image_file}!")
                except Exception as e:
                    st.error(f"Error during summarization for {image_file}: {e}")

    return st.session_state.summaries

def run_final_report(extracted_texts, summaries):
    st.header("Final Report Generation")
    segmented_images = sorted(os.listdir(SEGMENTED_DIR))

    data_map = {}
    for image_name in segmented_images:
        description = summaries.get(image_name, None)
        extracted_text = extracted_texts.get(image_name, None)
        data_map[image_name] = {
            "description": description,
            "extracted_text": extracted_text
        }

    json_path = os.path.join(REPORT_DIR, 'data_mapping.json')
    with open(json_path, 'w') as json_file:
        json.dump(data_map, json_file, indent=4)
    st.success(f"Data mapping saved to {json_path}")

    fig, ax = plt.subplots(figsize=(10, 10))
    original_image_path = os.path.join(INPUT_DIR, os.listdir(INPUT_DIR)[0])
    original_image = Image.open(original_image_path)
    ax.imshow(original_image)
    ax.set_title("Original Image with Segmented Object Annotations")

    for idx, image_name in enumerate(segmented_images):
        ax.annotate(f"Object {idx + 1}", xy=(10 + idx * 20, 20), color="red", fontsize=12)

    st.pyplot(fig)

    df_data_map = {
        "Object ID": [],
        "Description": [],
        "Extracted Text": [],
        "Image Name": []
    }

    for idx, image_name in enumerate(segmented_images):
        df_data_map["Object ID"].append(idx + 1)
        df_data_map["Description"].append(data_map[image_name]["description"])
        df_data_map["Extracted Text"].append(data_map[image_name]["extracted_text"])
        df_data_map["Image Name"].append(image_name)

    df = pd.DataFrame(df_data_map)
    st.dataframe(df)

# Main navigation logic
if page == "Upload Image":
    upload_image()
elif page == "Segment and Identify":
    run_segmentation_and_identification()
elif page == "Text Extraction":
    extracted_texts = run_text_extraction()
elif page == "Summarization":
    summaries = run_summarization()
elif page == "Final Report":
    if 'extracted_texts' in st.session_state and 'summaries' in st.session_state:
        run_final_report(st.session_state.extracted_texts, st.session_state.summaries)

# Footer
st.sidebar.markdown("### Footer")
st.sidebar.write("AI Pipeline Project for Image Analysis")
